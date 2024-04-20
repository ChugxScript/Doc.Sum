import spacy
from spacy.lang.en.stop_words import STOP_WORDS
from string import punctuation

# pip install python-docx
from docx import Document
from heapq import nlargest

# python -m spacy download en_core_web_sm

def summarize(file_content, summary_len):
    # create a list of stopwords.
    stopwords = list(STOP_WORDS)
    tokenize_doc = spacy.load('en_core_web_sm')

    # tokenizes the text, i.e. segments it into words, punctuation and so on
    file = tokenize_doc(file_content)
    tokens = [token.text for token in file]
    modified_punctuation = punctuation + '\n'

    word_freqs = word_frequency_table(file, stopwords, modified_punctuation)
    max_frequency = max(word_freqs.values())
    normalize_word_freqs = normalize_frequency(word_freqs, max_frequency)
    tokenized_sentence = [sent for sent in file.sents]
    sentence_scores = sentence_score(tokenized_sentence, normalize_word_freqs)
    summary_length = int(len(tokenized_sentence)*summary_len)
    init_summary = nlargest(summary_length, sentence_scores, key = sentence_scores.get)

    filter_summary = [word.text for word in init_summary]
    final_summary = ' '.join(filter_summary)
    # print("\nDoc:\n", file)
    # print("\n\nfinal_summary: \n", final_summary)
    
    # print("\nDoc len:", len(file))
    # print("\nfinal_summary len:", len(final_summary))

    return final_summary

def word_frequency_table(docz, stopwordsz, punctuationz):
    word_frequencies = {}
    for word in docz:
        if word.text.lower() not in stopwordsz:
            if word.text.lower() not in punctuationz:
                if word.text not in word_frequencies.keys():
                    word_frequencies[word.text] = 1
                else:
                    word_frequencies[word.text] += 1
    return word_frequencies

def normalize_frequency(word_freqsz, max_frequencyz):
    for word in word_freqsz.keys():
        word_freqsz[word] = word_freqsz[word]/max_frequencyz
    return word_freqsz

def sentence_score(tokenized_sentencez, normalize_word_freqsz):
    sentence_scores = {}
    for sent in tokenized_sentencez:
        for word in sent:
            if word.text.lower() in normalize_word_freqsz.keys():
                if sent not in sentence_scores.keys():
                    sentence_scores[sent] = normalize_word_freqsz[word.text.lower()]
                else:
                    sentence_scores[sent] += normalize_word_freqsz[word.text.lower()]
    return sentence_scores




# this code will be remove since we read the file in the views.py
# DOCDOCX = './Doc.docx'
def read_document(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    print("ORIGINAL TEXT:")
    print(text)
    print()
    return text
# summarize(read_document(DOCDOCX))













# import spacy
# from spacy.lang.en.stop_words import STOP_WORDS
# from string import punctuation

# # create a list of stopwords.
# stopwords = list(STOP_WORDS)
# tokenize_doc = spacy.load('en_core_web_sm')

# DOCDOCX = './Doc.docx'
# from docx import Document
# def read_document(file_path):
#     doc = Document(file_path)
#     text = ""
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + "\n"
    
#     print("ORIGINAL TEXT:")
#     print(text)
#     print()
#     return text

# # tokenizes the text, i.e. segments it into words, punctuation and so on
# doc = tokenize_doc(read_document(DOCDOCX))

# tokens = [token.text for token in doc]
# # print(tokens)

# punctuation = punctuation + '\n'


# def word_frequency_table(docz):
#     word_frequencies = {}
#     for word in docz:
#         if word.text.lower() not in stopwords:
#             if word.text.lower() not in punctuation:
#                 if word.text not in word_frequencies.keys():
#                     word_frequencies[word.text] = 1
#                 else:
#                     word_frequencies[word.text] += 1
#     return word_frequencies

# word_freqs = word_frequency_table(doc)
# # print(word_freqs)
# max_frequency = max(word_freqs.values())
# # print(max_frequency)

# def normalize_frequency(word_freqsz, max_frequencyz):
#     for word in word_freqsz.keys():
#         word_freqsz[word] = word_freqsz[word]/max_frequencyz
#     return word_freqsz
# normalize_word_freqs = normalize_frequency(word_freqs, max_frequency)
# # print(normalize_word_freqs)

# def tokenize_sentence(docz):
#     return [sent for sent in docz.sents]
# tokenized_sentence = tokenize_sentence(doc)
# # print(tokenized_sentence)

# def sentence_score(tokenized_sentencez, normalize_word_freqsz):
#     sentence_scores = {}
#     for sent in tokenized_sentencez:
#         for word in sent:
#             if word.text.lower() in normalize_word_freqsz.keys():
#                 if sent not in sentence_scores.keys():
#                     sentence_scores[sent] = normalize_word_freqsz[word.text.lower()]
#                 else:
#                     sentence_scores[sent] += normalize_word_freqsz[word.text.lower()]
#     return sentence_scores
# sentence_scores = sentence_score(tokenized_sentence, normalize_word_freqs)
# # print(sentence_scores)

# from heapq import nlargest
# summary_length = int(len(tokenized_sentence)*0.3)

# summary = nlargest(summary_length, sentence_scores, key = sentence_scores.get)
# # print(summary)

# final_summary = [word.text for word in summary]
# summaryz = ' '.join(final_summary)

# print("\n\nsummaryz: ", summaryz)